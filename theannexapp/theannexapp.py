import openpyxl
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.shared import Mm
import streamlit as st
from io import BytesIO
import pandas as pd
import numpy as np
import datetime


#------ Funcion auxiliar ---------

def set_table_font_size(table, size_pt):
    """
    Cambia el tamaño de fuente de todo el texto en una tabla.
    :param table: objeto Table de python-docx
    :param size_pt: tamaño en puntos (int o float)
    """
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(size_pt)

def armar_esqueleto(documento, planilla, oficina):
    """
    Arma la hoja con determinado formato del word.
    
    :param documento: Documento que se está escribiendo
    :param planilla: Planilla actual de la cuál se sacan los datos
    :param oficina: [int, str], array con el número de la oficina y el nombre de la oicina
    """

    parrafo_exp = documento.add_paragraph()
    # El nombre del anexo correspondiente a la oficina es del tipo "oficina - nombre de la oficina"
    nombre_anexo = str(oficina[0]) + " - " + oficina[1]
    run_exp = parrafo_exp.add_run(nombre_anexo)
    run_exp.bold = True
    run_exp.underline = True
    run_exp.font.size = Pt(16)
    parrafo_exp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Tabla
    tabla = documento.add_table(rows=1, cols=9)
    tabla.style = 'Table Grid'  # Bordes visibles

    encabezado = tabla.rows[0].cells
    encabezado[0].text = "NRO.OFICINA"
    encabezado[1].text = "LEGAJO"
    encabezado[2].text = "APELLIDO Y NOMBRE"
    encabezado[3].text = "CATEGORÍA"
    encabezado[4].text = "FUNCIÓN"
    encabezado[5].text = "BONIFICACIÓN"
    encabezado[6].text = "INGRESO"
    encabezado[7].text = "EGRESO"
    encabezado[8].text = "NOTIFICACION FIRMA Y FECHA"

    return tabla

def armar_anexo(documento,planilla):
    """
    Pasa los datos de la planilla a un formato tabla en un Word.
    
    :param documento: Documento que se está escribiendo.
    :param planilla: Planilla  .xlsx de la cuál se están sacando los datos.
    """
    
    wb = openpyxl.load_workbook(planilla,read_only = True)
    ws = wb.worksheets[0]
    
    oficina_anterior = ws.cell(row=2, column=1).value #Primer número de oficina del área
    numero_oficina = ws.cell(row=2, column=1).value
    nombre_oficina = ws.cell(row=2, column=2).value
    tabla = armar_esqueleto(documento, planilla,[numero_oficina, nombre_oficina])

    for row in ws.iter_rows(min_row = 2, max_row = ws.max_row, min_col = 1, max_col = 10):

        row_sin_col2 = [] # Para sacar la columna con el nombre de la oficina

        for i, cell in enumerate(row):
            if i == 1:
                nombre_oficina = cell.value #Guardo el nombre de la oficina
                continue
            else:
                row_sin_col2.append(cell.value)

        oficina = row_sin_col2[0] 
        numero_oficina = row_sin_col2[0] #Guardo el numero de oficina


        if oficina != oficina_anterior:
            oficina_anterior = oficina

            documento.add_page_break()
            tabla = armar_esqueleto(documento, planilla,[numero_oficina,nombre_oficina])
            
        if not all(cell is None for cell in row_sin_col2): # SIN la columna del nombre de la oficina
            fila = tabla.add_row().cells
            
            for i, cell in enumerate(row_sin_col2):
                if i in [6, 7] and cell is not None:
                    fila[i].text = cell.strftime("%d/%m/%Y")
                elif cell is not None:
                    fila[i].text = str(cell)
                else:
                    fila[i].text = ""
            # espacio para firmar
            if len(fila) > 8:
                fila[8].text = ""

        set_table_font_size(tabla, 10)

    documento.add_page_break()

def armar_anexos(planillas):
    """
    Para cada archivo subido se escribe en el documento word lo requerido
    
    :param planillas: planillas subidas
    """
    documento = Document()

    style = documento.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

    section = documento.sections[0]
    section.page_height = Mm(210)
    section.page_width = Mm(297)
    section.left_margin = Mm(25.4)
    section.right_margin = Mm(25.4)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)
    section.orientation = WD_ORIENT.LANDSCAPE

    for planilla in planillas:
        armar_anexo(documento,planilla)

    return documento

## Streamlit APP ## 

st.title('The Annex App📎')
st_archivos = st.file_uploader("Clickeá donde dice 'Browse files' y subí los archivos", accept_multiple_files=True)

if st_archivos:
    st.success(f"Subiste {len(st_archivos)} planilla(s)")
    titulo = st.text_input("Escribí el nombre del archivo y presioná Enter", "Anexo Subsecretaría ABC")

    if st.button("Procesar y armar anexos"):
        documento = armar_anexos(st_archivos)

        buffer = BytesIO()
        documento.save(buffer)
        buffer.seek(0)

        st.info('Recordá revisar el documento')
        st.download_button(
            label="Descargar notificaciones",
            data=buffer,
            file_name= titulo.strip() + ".docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            icon=":material/download:",
        )